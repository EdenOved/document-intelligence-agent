from __future__ import annotations

from io import BytesIO
from pathlib import Path

import fitz
from docx import Document as DocxDocument
from fastapi import UploadFile

from app.schemas import ParsedDocument, ParsedPage


async def parse_upload_file(file: UploadFile) -> ParsedDocument:
    filename = file.filename or "uploaded_file"
    extension = Path(filename).suffix.lower()

    content = await file.read()
    if not content:
        raise RuntimeError("Uploaded file is empty")

    if extension == ".pdf":
        return _parse_pdf_bytes(filename, content)

    if extension == ".docx":
        return _parse_docx_bytes(filename, content)

    if extension == ".txt":
        return _parse_txt_bytes(filename, content)

    raise RuntimeError(f"Unsupported file type: {extension}")


def _parse_pdf_bytes(filename: str, content: bytes) -> ParsedDocument:
    pages: list[ParsedPage] = []

    with fitz.open(stream=content, filetype="pdf") as pdf:
        for page_index, page in enumerate(pdf, start=1):
            text = page.get_text("text") or ""
            pages.append(
                ParsedPage(
                    page_number=page_index,
                    text=_normalize_text(text),
                    metadata={"page_number": page_index},
                )
            )

    return ParsedDocument(
        filename=filename,
        source_type="pdf",
        pages=pages,
        metadata={"source_extension": ".pdf"},
    )


def _parse_docx_bytes(filename: str, content: bytes) -> ParsedDocument:
    doc = DocxDocument(BytesIO(content))

    paragraphs = []
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            paragraphs.append(text)

    full_text = "\n".join(paragraphs)

    return ParsedDocument(
        filename=filename,
        source_type="docx",
        pages=[
            ParsedPage(
                page_number=1,
                text=_normalize_text(full_text),
                metadata={"page_number": 1},
            )
        ],
        metadata={"source_extension": ".docx"},
    )


def _parse_txt_bytes(filename: str, content: bytes) -> ParsedDocument:
    text = content.decode("utf-8", errors="ignore")

    return ParsedDocument(
        filename=filename,
        source_type="txt",
        pages=[
            ParsedPage(
                page_number=1,
                text=_normalize_text(text),
                metadata={"page_number": 1},
            )
        ],
        metadata={"source_extension": ".txt"},
    )


def _normalize_text(text: str) -> str:
    lines = [line.rstrip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    cleaned_lines = []
    previous_blank = False

    for line in lines:
        stripped = line.strip()

        if not stripped:
            if not previous_blank:
                cleaned_lines.append("")
            previous_blank = True
            continue

        cleaned_lines.append(stripped)
        previous_blank = False

    return "\n".join(cleaned_lines).strip()